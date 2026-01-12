#!/usr/bin/env python3
import os
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import json

def extract_content_and_images(docx_path):
    doc = Document(docx_path)

    content = {
        'paragraphs': [],
        'tables': [],
        'images': []
    }

    # Extract images
    image_dir = 'extracted_images'
    os.makedirs(image_dir, exist_ok=True)

    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_count += 1
            image = rel.target_part.blob
            image_ext = rel.target_ref.split('.')[-1]
            image_filename = f'image_{image_count}.{image_ext}'
            image_path = os.path.join(image_dir, image_filename)

            with open(image_path, 'wb') as img_file:
                img_file.write(image)

            content['images'].append({
                'filename': image_filename,
                'path': image_path
            })

    # Extract paragraphs with style information
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            para_info = {
                'text': paragraph.text,
                'style': paragraph.style.name if paragraph.style else 'Normal',
                'bold': any(run.bold for run in paragraph.runs),
                'italic': any(run.italic for run in paragraph.runs)
            }
            content['paragraphs'].append(para_info)

    # Extract tables
    for table_idx, table in enumerate(doc.tables):
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            table_data.append(row_data)
        content['tables'].append({
            'index': table_idx,
            'data': table_data
        })

    return content

if __name__ == '__main__':
    content = extract_content_and_images('Profil Lembaga TK ABA Kalikajar New.docx')

    # Save to JSON
    with open('extracted_content.json', 'w', encoding='utf-8') as f:
        json.dump(content, f, ensure_ascii=False, indent=2)

    # Also create a readable markdown file
    with open('extracted_content.md', 'w', encoding='utf-8') as f:
        f.write('# TK ABA Kalikajar - Extracted Content\n\n')

        for para in content['paragraphs']:
            if 'Heading' in para['style']:
                level = '##' if '1' in para['style'] else '###'
                f.write(f"\n{level} {para['text']}\n\n")
            else:
                f.write(f"{para['text']}\n\n")

        if content['tables']:
            f.write('\n## Tables\n\n')
            for table in content['tables']:
                f.write(f"\n### Table {table['index'] + 1}\n\n")
                for row in table['data']:
                    f.write('| ' + ' | '.join(row) + ' |\n')
                f.write('\n')

        if content['images']:
            f.write('\n## Images\n\n')
            for img in content['images']:
                f.write(f"- {img['filename']}\n")

    print(f"Extraction complete!")
    print(f"- Found {len(content['paragraphs'])} paragraphs")
    print(f"- Found {len(content['tables'])} tables")
    print(f"- Found {len(content['images'])} images")
